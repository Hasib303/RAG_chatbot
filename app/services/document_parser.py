from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.services.domain import TextSection


def parse_document(file_path: Path) -> list[TextSection]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    if suffix == ".docx":
        return _parse_docx(file_path)
    raise ValueError("Only PDF and DOCX files are supported.")


def _parse_pdf(file_path: Path) -> list[TextSection]:
    reader = PdfReader(str(file_path))
    sections: list[TextSection] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = _normalise_whitespace(page.extract_text() or "")
        if not text:
            continue
        sections.append(
            TextSection(
                text=text,
                source_label=f"Page {page_number}",
                page_number=page_number,
            )
        )

    if not sections:
        raise ValueError("The PDF did not contain extractable text.")
    return sections


def _parse_docx(file_path: Path) -> list[TextSection]:
    document = DocxDocument(str(file_path))
    paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = _normalise_whitespace(paragraph.text)
        if text:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [_normalise_whitespace(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError("The DOCX did not contain extractable text.")

    group_size = 6
    sections: list[TextSection] = []
    for index in range(0, len(paragraphs), group_size):
        group = paragraphs[index : index + group_size]
        start = index + 1
        end = index + len(group)
        sections.append(
            TextSection(
                text="\n".join(group),
                source_label=f"Paragraphs {start}-{end}",
            )
        )

    return sections


def _normalise_whitespace(text: str) -> str:
    return " ".join(text.split())
